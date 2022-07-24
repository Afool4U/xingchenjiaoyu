import winreg

from PyPDF2 import PdfFileReader, PdfFileWriter
from docx2pdf import convert
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入库：设置对象居中、对齐等
from docx.oxml import OxmlElement, ns
from win32com import client as wc

cwd = os.getcwd()


subject = None


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.text = ''
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    t1.text = '第'
    page_run._r.append(t1)
    page_num_run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    instrText.text = "PAGE"
    create_attribute(instrText, 'w:fldCharType', 'separate')
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    t2.text = '页，共'
    of_run._r.append(t2)
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')
    instrText2 = create_element('w:instrText')
    instrText2.text = "NUMPAGES"
    create_attribute(instrText2, 'w:fldCharType', 'separate')
    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')
    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)
    of_run = paragraph.add_run()
    t3 = create_element('w:t')
    t3.text = '页'
    of_run._r.append(t3)


def docToDocx(fileName):
    # print("开始处理     文件名：" + fileName)
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(fileName)
    # [:-4]的意思是选这个字符串从开始到最后倒数第4位（不含）
    docxNamePath = fileName + 'x'
    # print('转换完成！' + docxNamePath)
    doc.SaveAs(docxNamePath, 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()
    # 一定要记得关闭docx，否则会出现文件占用
    os.remove(fileName)
    return fileName + 'x'


def get_desktop():
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    return winreg.QueryValueEx(key, "Desktop")[0]


def replaceHeader(sFileName):
    if sFileName.endswith(".docx") and ("~$" not in sFileName):
        # 形成docx对象
        document = Document(sFileName)
        header = document.sections[0].header
        paragraph = header.paragraphs[0]
        paragraph.text = "星辰2022" + subject + "更新版"
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置页眉居中对齐
        document.sections[0].footer.is_linked_to_previous = True
        add_page_number(document.sections[0].footer.paragraphs[0])

        document.save(sFileName)
    else:
        pass
        # print("Not doc File")
    # print('{}页眉页脚添加成功！'.format(sFileName))


all_docs = []


def get_all_files(path=os.path.join(os.getcwd(), '待处理文件')):
    file_list = os.listdir(path)
    for file in file_list:
        if os.path.isdir(os.path.join(path, file)):
            get_all_files(os.path.join(path, file))
        elif file.endswith(('.doc', '.docx')) and not file.startswith('~'):  # 排除临时文件
            all_docs.append(os.path.join(path, file))


def add_watermark(pdf_file_in, pdf_file_mark, pdf_file_out):
    """把水印添加到pdf中"""

    pdf_output = PdfFileWriter()
    input_stream = open(pdf_file_in, 'rb')
    pdf_input = PdfFileReader(input_stream, strict=False)
    # 获取PDF文件的页数
    pageNum = pdf_input.getNumPages()
    # 读入水印pdf文件
    pdf_watermark = PdfFileReader(open(pdf_file_mark, 'rb'), strict=False)
    # 给每一页打水印
    for i in range(pageNum):
        page = pdf_input.getPage(i)
        page.mergePage(pdf_watermark.getPage(0))
        page.compressContentStreams()  # 压缩内容
        pdf_output.addPage(page)
        pdf_output.write(open(pdf_file_out, 'wb'))
    input_stream.close()
    # print('{}水印添加成功！'.format(pdf_file_out))


def docx2pdf(file_path):
    # doc 文件 需要先转为 docx 文件（一般情况下、直接修改后缀名、不会对原文件有影响）
    inputFile = file_path  # 要转换的文件：已存在
    suffix = '.docx'
    prefix = file_path[0:file_path.rfind(suffix)]
    outputFile = prefix + ".pdf"
    convert(inputFile, outputFile)


def pdf_encrypt(pdf_file_path):
    # print('{}加密成功！'.format(pdf_file_path))
    pdf_reader = PdfFileReader(pdf_file_path)
    pdf_writer = PdfFileWriter()

    for page in range(pdf_reader.getNumPages()):
        pdf_writer.addPage(pdf_reader.getPage(page))
    pdf_writer.encrypt(user_pwd='', owner_pwd="775828", permissions_flag=-3900)  # 设置密码
    pdf_temp_path = os.path.join(pdf_file_path[0:pdf_file_path.rfind('\\')], 'temp.pdf')
    if os.path.exists(pdf_temp_path):
        os.remove(pdf_temp_path)
    with open(pdf_temp_path, 'wb') as out:
        pdf_writer.write(out)


"""
if __name__ == '__main__':
    get_all_files(cwd)
    error_files = []
    total_num = len(all_docs)
    for idx, docx in enumerate(all_docs):
        try:
            print('处理进度:{}/{} '.format(idx + 1, total_num), end='')
            if docx.endswith('.doc'):
                docx = docToDocx(docx)
            print('正在处理' + docx[docx.rfind('\\') + 1:] + '文件！')
            # 给word加页眉和页脚
            replaceHeader(docx)
            # 处理后的word转为pdf
            docx2pdf(docx)
            suffix = '.docx'
            prefix = docx[0:docx.rfind(suffix)]
            pdf_file_path = prefix + ".pdf"
            # 给转换好的pdf加水印
            pdf_temp_path = os.path.join(docx[0:docx.rfind('\\')], 'temp.pdf')
            if os.path.exists(pdf_temp_path):
                os.remove(pdf_temp_path)
            add_watermark(pdf_file_path, os.path.join(cwd, r'resource\watermark.pdf'), pdf_temp_path)
            os.remove(pdf_file_path)
            os.rename(pdf_temp_path, pdf_file_path)
            pdf_encrypt(pdf_file_path)
            os.remove(pdf_file_path)
            os.remove(docx)
            os.rename(pdf_temp_path, pdf_file_path)
        except Exception as e:
            error_files.append(docx)
            print(e)
    if len(error_files) != 0:
        print('\n转换失败的文件：')
        for f in error_files:
            print(f)
    else:
        print('\n全部转换成功！')
    input()
    """


def replace_placeholder(title, template_docx='head.docx'):
    doc = Document(template_docx)
    """替换占位符"""
    params = {
        "一元二次方程": title,
    }
    for paragraph in doc.paragraphs:
        for param in params:
            pv = params[param]
            ph = f'ph_{param}'
            if ph in paragraph.text:
                for run in paragraph.runs:
                    if ph in run.text:
                        run.text = run.text.replace(ph, pv)
    doc.save('result.docx')


def find_and_del_title(docx_file):
    doc = Document(docx_file)
    # print(docx_file)
    paragraph = doc.paragraphs[0]
    for run in paragraph.runs:
        run.text = run.text.replace(run.text, '')
    doc.save(docx_file)
    return docx_file.split('《')[1].split("》")[0]


def merge_docx(docx_file_list):  # 绝对路径
    word = wc.Dispatch('Word.Application')
    word.Visible = False
    output = word.Documents.Add()
    # 新建空的word文档,
    for file in docx_file_list:
        output.Application.Selection.InsertFile(file)
    # 拼接文档
    doc = output.Range(output.Content.Start, output.Content.End)
    # 获取合并后文档的内容
    output.SaveAs(os.getcwd() + '\\output.docx')  # 保存
    output.Close()  # 关闭
    word.Quit()


def add_head(docx_file):
    title = find_and_del_title(docx_file)
    print('title: ', title[title.find(' ') + 1:])
    replace_placeholder(title[title.find(' ') + 1:])
    merge_docx([os.getcwd() + '\\result.docx', docx_file])
    os.remove(docx_file)
    os.rename(os.getcwd() + '\\output.docx', docx_file)


if __name__ == '__main__':
    get_all_files()
    error_files = []
    total_num = len(all_docs)
    for idx, docx in enumerate(all_docs):
        try:
            print('处理进度:{}/{} '.format(idx + 1, total_num), end='')
            if docx.endswith('.doc'):
                docx = docToDocx(docx)
            print('正在处理' + docx[docx.rfind('\\') + 1:] + '文件！')
            subject = '数学' if '数学' in docx else '化学'  # 自动识别学科
            add_head(docx)
            # 给word加页眉和页脚
            replaceHeader(docx)
            # 处理后的word转为pdf
            docx2pdf(docx)
            suffix = '.docx'
            prefix = docx[0:docx.rfind(suffix)]
            pdf_file_path = prefix + ".pdf"
            # 给转换好的pdf加水印
            pdf_temp_path = os.path.join(docx[0:docx.rfind('\\')], 'temp.pdf')
            if os.path.exists(pdf_temp_path):
                os.remove(pdf_temp_path)
            add_watermark(pdf_file_path, os.path.join(cwd, r'resource\watermark.pdf'), pdf_temp_path)
            os.remove(pdf_file_path)
            os.rename(pdf_temp_path, pdf_file_path)
            pdf_encrypt(pdf_file_path)
            os.remove(pdf_file_path)
            os.remove(docx)
            os.rename(pdf_temp_path, pdf_file_path)
        except Exception as e:
            error_files.append(docx)
            print(e)
    if len(error_files) != 0:
        print('\n转换失败的文件：')
        for f in error_files:
            print(f)
    else:
        print('\n全部转换成功！')
    input()

